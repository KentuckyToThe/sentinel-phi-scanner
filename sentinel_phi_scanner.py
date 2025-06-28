"""
🔐 Sentinel: AI PHI Risk Scanner
A Streamlit app for detecting Protected Health Information (PHI) in documents

Required packages:
pip install streamlit pandas PyPDF2 python-docx

Run with:
streamlit run sentinel_phi_scanner.py
"""

import streamlit as st
import pandas as pd
import re
from datetime import datetime
import io
import base64
import json
from typing import Dict, List, Tuple
import time

# Import libraries for file processing
try:
    import PyPDF2
except ImportError:
    st.error("Please install PyPDF2: pip install PyPDF2")
    st.stop()

try:
    import docx
except ImportError:
    docx = None  # Optional for Word docs

# Page configuration
st.set_page_config(
    page_title="Sentinel: PHI Risk Scanner",
    page_icon="🔐",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main {
        padding-top: 2rem;
    }
    .stButton>button {
        width: 100%;
        border-radius: 8px;
        height: 3em;
        font-weight: 600;
    }
    .risk-high {
        background-color: #ff4757;
        color: white;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-weight: bold;
    }
    .risk-moderate {
        background-color: #ffa502;
        color: white;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-weight: bold;
    }
    .risk-low {
        background-color: #2ed573;
        color: white;
        padding: 0.25rem 0.75rem;
        border-radius: 20px;
        font-weight: bold;
    }
    div.stAlert {
        border-radius: 10px;
    }
    .stDownloadButton {
        display: inline-block;
    }
    .risk-card {
        border-radius: 10px;
        padding: 20px;
        margin: 20px 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    .risk-high-card {
        background: linear-gradient(135deg, #ffebee 0%, #ffcdd2 100%);
        border-left: 5px solid #ff4757;
    }
    .risk-moderate-card {
        background: linear-gradient(135deg, #fff8e1 0%, #ffe0b2 100%);
        border-left: 5px solid #ff9f43;
    }
    .risk-low-card {
        background: linear-gradient(135deg, #e8f5e9 0%, #c8e6c9 100%);
        border-left: 5px solid #10ac84;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'scan_history' not in st.session_state:
    st.session_state.scan_history = []
if 'total_files_scanned' not in st.session_state:
    st.session_state.total_files_scanned = 0
if 'high_risk_count' not in st.session_state:
    st.session_state.high_risk_count = 0
if 'total_phi_found' not in st.session_state:
    st.session_state.total_phi_found = 0
if 'current_results' not in st.session_state:
    st.session_state.current_results = []
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = 0

# PHI Detection Patterns - Advanced Version
PHI_PATTERNS = {
    'SSN': {
        'pattern': r'\b(?:SSN|Social Security|SS#|SS\s*#|Social\s*Security\s*Number)[\s:]*(\d{3}-\d{2}-\d{4}|\d{9})\b',
        'description': 'Social Security Number',
        'priority': 'high'
    },
    'MRN': {
        'pattern': r'\b(?:MRN|Medical Record|MR#|Account#|Acct#|Patient ID)[\s:#]*([A-Z]{0,3}\d{6,12})\b',
        'description': 'Medical Record Number',
        'priority': 'high'
    },
    'DOB': {
        'pattern': r'\b(?:DOB|Date of Birth|Birth Date|Born|Birthdate)[\s:]*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4}|\d{1,2}[-/]\d{1,2}[-/]\d{2})\b',
        'description': 'Date of Birth',
        'priority': 'medium'
    },
    'Phone': {
        'pattern': r'\b(?:Phone|Ph|Tel|Telephone|Cell|Mobile|Contact)[\s:#]*\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b',
        'description': 'Phone Number',
        'priority': 'medium'
    },
    'Email': {
        'pattern': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        'description': 'Email Address',
        'priority': 'low'
    },
    'Patient_Name': {
        'pattern': r'(?:Patient Name|Patient|Name|Pt)[\s:]+([A-Z][a-z]+(?:\s+[A-Z]\.?)?\s+[A-Z][a-z]+)',
        'description': 'Patient Name',
        'priority': 'high'
    },
    'Address': {
        'pattern': r'\b\d{1,5}\s+(?:[NSEW]\.?\s+)?(?:\w+\s+){0,2}(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Circle|Cir|Plaza|Pl|Way|Parkway|Pkwy)(?:\s+(?:Apt|Apartment|Suite|Ste|Unit|#)\s*\w+)?\b',
        'description': 'Physical Address',
        'priority': 'medium'
    },
    'Credit_Card': {
        'pattern': r'\b(?:4[0-9]{12}(?:[0-9]{3})?|5[1-5][0-9]{14}|3[47][0-9]{13}|3(?:0[0-5]|[68][0-9])[0-9]{11}|6(?:011|5[0-9]{2})[0-9]{12})\b',
        'description': 'Credit Card Number',
        'priority': 'high'
    },
    'Insurance_ID': {
        'pattern': r'\b(?:Insurance ID|Member ID|Policy#|Group#)[\s:]*([A-Z0-9]{8,15})\b',
        'description': 'Insurance ID',
        'priority': 'medium'
    },
    'License': {
        'pattern': r'\b(?:DL|Driver\'s License|License#)[\s:]*([A-Z]{1,2}\d{5,8})\b',
        'description': 'Driver License',
        'priority': 'medium'
    }
}

# Medical terms to exclude from name detection
MEDICAL_TERMS = {
    'medical', 'history', 'physical', 'examination', 'emergency', 'department',
    'chief', 'complaint', 'present', 'illness', 'review', 'systems', 'vital', 'signs',
    'blood', 'pressure', 'pulse', 'temperature', 'general', 'appearance', 'skin',
    'chest', 'abdomen', 'extremities', 'neurological', 'assessment', 'plan',
    'diagnosis', 'treatment', 'medication', 'allergy', 'social', 'family',
    'surgical', 'medical', 'problem', 'list', 'discharge', 'summary',
    'heart', 'lung', 'liver', 'kidney', 'brain', 'spine', 'bone',
    'normal', 'abnormal', 'positive', 'negative', 'stable', 'acute', 'chronic',
    'bilateral', 'left', 'right', 'upper', 'lower', 'anterior', 'posterior'
}

# Provider indicators - names with these are likely healthcare providers, not patients
PROVIDER_INDICATORS = {
    'dr', 'doctor', 'physician', 'md', 'do', 'phd', 'np', 'pa', 'rn', 'nurse',
    'practitioner', 'therapist', 'specialist', 'attending', 'resident', 'intern',
    'professor', 'chief', 'director', 'administrator', 'staff', 'provider',
    'surgeon', 'radiologist', 'pathologist', 'anesthesiologist', 'psychiatrist',
    'cardiologist', 'neurologist', 'oncologist', 'pediatrician', 'obstetrician'
}

# Context indicators that strongly suggest a patient name
PATIENT_CONTEXT_INDICATORS = {
    'patient', 'pt', 'client', 'member', 'beneficiary', 'enrollee', 'applicant',
    'claimant', 'subject', 'individual', 'person', 'admitted', 'discharged',
    'transferred', 'presenting', 'complaining', 'suffering', 'diagnosed',
    'treated', 'scheduled', 'underwent', 'received', 'declined', 'refused'
}

# Common non-patient name contexts
NON_PATIENT_CONTEXTS = {
    'authored', 'written', 'prepared', 'reviewed', 'approved', 'signed',
    'performed', 'conducted', 'administered', 'supervised', 'consulted',
    'referred', 'ordered', 'prescribed', 'documented', 'reported',
    'hospital', 'clinic', 'center', 'institute', 'university', 'college',
    'department', 'unit', 'ward', 'office', 'laboratory', 'pharmacy'
}

def is_medical_term(text: str) -> bool:
    """Check if text contains common medical terms"""
    words = text.lower().split()
    return any(word in MEDICAL_TERMS for word in words)

def analyze_name_context(text: str, start_pos: int, end_pos: int, window_size: int = 100) -> Dict[str, any]:
    """Analyze the context around a detected name to determine if it's a patient name"""
    # Get surrounding context
    context_start = max(0, start_pos - window_size)
    context_end = min(len(text), end_pos + window_size)
    context = text[context_start:context_end].lower()
    
    # Get the actual name for analysis
    name_text = text[start_pos:end_pos]
    
    # Initialize scores
    patient_score = 0
    provider_score = 0
    confidence = 50  # Base confidence
    
    # Check for provider indicators
    for indicator in PROVIDER_INDICATORS:
        if indicator in context:
            provider_score += 10
            if f"{indicator}." in context or f"{indicator} " in context:
                provider_score += 5
    
    # Check for patient context indicators
    for indicator in PATIENT_CONTEXT_INDICATORS:
        if indicator in context:
            patient_score += 10
            # Strong indicators get extra weight
            if indicator in ['patient', 'pt', 'admitted', 'diagnosed', 'treated']:
                patient_score += 10
    
    # Check for non-patient contexts
    for indicator in NON_PATIENT_CONTEXTS:
        if indicator in context:
            provider_score += 5
    
    # Check specific patterns
    if re.search(r'(patient|pt\.?)\s*(name|:|\s)\s*' + re.escape(name_text), context, re.IGNORECASE):
        patient_score += 30
        confidence += 20
    
    if re.search(r'(doctor|dr\.?|physician)\s+' + re.escape(name_text), context, re.IGNORECASE):
        provider_score += 30
        confidence += 20
    
    # Check for medical credentials after name
    if re.search(re.escape(name_text) + r'\s*,?\s*(md|do|phd|rn|np|pa)', context, re.IGNORECASE):
        provider_score += 40
    
    # Check for dates near the name (often indicates patient context)
    date_pattern = r'\b(?:DOB|born|birth\s*date|date\s*of\s*birth)[:\s]*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}'
    if re.search(date_pattern, context, re.IGNORECASE):
        patient_score += 20
    
    # Check for MRN or account numbers near the name
    if re.search(r'(mrn|medical\s*record|account)\s*#?\s*:?\s*\w+', context, re.IGNORECASE):
        patient_score += 20
    
    # Determine if it's likely a patient name
    is_likely_patient = patient_score > provider_score
    
    # Calculate final confidence
    if is_likely_patient:
        confidence = min(95, confidence + (patient_score - provider_score))
    else:
        confidence = max(20, confidence - (provider_score - patient_score))
    
    # Adjust priority based on context
    if is_likely_patient and patient_score > 30:
        priority = 'high'
    elif is_likely_patient and patient_score > 15:
        priority = 'medium'
    else:
        priority = 'low'
    
    return {
        'is_patient': is_likely_patient,
        'confidence': confidence,
        'patient_score': patient_score,
        'provider_score': provider_score,
        'priority': priority,
        'context_type': 'patient' if is_likely_patient else 'provider/other'
    }

def smart_ssn_detection(text: str, match: re.Match) -> Dict[str, any]:
    """Smart SSN detection with context analysis"""
    value = match.group(1) if match.groups() else match.group(0)
    context = text[max(0, match.start()-50):min(len(text), match.end()+50)]
    
    # Check if it's actually an SSN or might be something else
    confidence = 70
    priority = 'medium'
    
    # High confidence patterns
    if re.search(r'(social\s*security|ssn|ss\s*#)', context, re.IGNORECASE):
        confidence = 95
        priority = 'high'
    # Could be a phone number or other ID
    elif re.search(r'(phone|tel|fax|id|account|reference)', context, re.IGNORECASE):
        confidence = 40
        priority = 'low'
    # Format check
    elif re.match(r'^\d{3}-\d{2}-\d{4}

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Initialize default values
            confidence = 70
            priority = config.get('priority', 'medium')
            skip_finding = False
            
            # Special handling for patient names with advanced context analysis
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
                
                # Analyze context to determine if it's really a patient name
                context_analysis = analyze_name_context(text, match.start(), match.end())
                
                if not context_analysis['is_patient']:
                    # Skip non-patient names unless confidence is very low
                    if context_analysis['confidence'] < 30:
                        continue
                
                confidence = context_analysis['confidence']
                priority = context_analysis['priority']
                
                # Add context type to description
                config['description'] = f"{config['description']} ({context_analysis['context_type']})"
            
            # Smart SSN detection
            elif phi_type == 'SSN':
                ssn_analysis = smart_ssn_detection(text, match)
                confidence = ssn_analysis['confidence']
                priority = ssn_analysis['priority']
            
            # Enhanced DOB detection
            elif phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Separate findings by actual risk level (using the smart priority from detection)
    high_risk_items = [f for f in findings if f.get('priority') == 'high' and f.get('confidence', 0) >= 70]
    medium_risk_items = [f for f in findings if f.get('priority') == 'medium' and f.get('confidence', 0) >= 60]
    low_risk_items = [f for f in findings if f.get('priority') == 'low']
    
    # Count unique PHI types at each risk level
    high_risk_types = set(f['type'] for f in high_risk_items)
    medium_risk_types = set(f['type'] for f in medium_risk_items)
    
    # Calculate base score with smart weighting
    high_risk_score = len(high_risk_items) * 15 + len(high_risk_types) * 10
    medium_risk_score = len(medium_risk_items) * 5 + len(medium_risk_types) * 5
    low_risk_score = len(low_risk_items) * 2
    
    # Bonus for particularly sensitive combinations
    sensitive_types = {'SSN', 'Credit Card', 'MRN'}
    if len(high_risk_types.intersection(sensitive_types)) >= 2:
        high_risk_score += 20
    
    # Check for identity theft risk (name + SSN or name + DOB + address)
    has_patient_name = any(f['type'] == 'Patient Name' and f.get('priority') == 'high' for f in findings)
    has_ssn = any(f['type'] == 'SSN' and f.get('confidence', 0) >= 80 for f in findings)
    has_dob = any(f['type'] == 'DOB' for f in findings)
    has_address = any(f['type'] == 'Address' for f in findings)
    
    if has_patient_name and has_ssn:
        high_risk_score += 25
    elif has_patient_name and has_dob and has_address:
        high_risk_score += 20
    
    # Total score calculation
    total_score = high_risk_score + medium_risk_score + low_risk_score
    
    # Determine risk level with more nuanced thresholds
    if len(high_risk_items) >= 3 or total_score >= 80 or (has_patient_name and has_ssn):
        return 'high', min(total_score, 95)
    elif len(high_risk_items) >= 1 or len(medium_risk_items) >= 3 or total_score >= 40:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Special handling for patient names
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
            
            # Skip dates that are clearly not DOB (like years only)
            if phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})$')
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                except:
                    pass
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': calculate_confidence(phi_type, value, full_match),
                'priority': config.get('priority', 'medium')
            }
            
            # Only add if confidence is above threshold
            if finding['confidence'] >= 60:  # Minimum 60% confidence
                findings.append(finding)
                found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
)
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                        
                        # Check context for DOB indicators
                        context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                        if any(term in context for term in ['dob', 'birth', 'born']):
                            confidence = 90
                            priority = 'medium'
                        else:
                            confidence = 50
                            priority = 'low'
                except:
                    confidence = 60
            
            # Enhanced MRN detection
            elif phi_type == 'MRN':
                context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                if any(term in context for term in ['mrn', 'medical record', 'mr#']):
                    confidence = 90
                    priority = 'high'
                elif any(term in context for term in ['account', 'acct']):
                    confidence = 75
                    priority = 'medium'
                else:
                    confidence = 50
                    priority = 'low'
            
            # Phone number validation
            elif phi_type == 'Phone':
                # Check if it's actually a phone number format
                phone_value = re.sub(r'[^\d]', '', value)
                if len(phone_value) == 10:
                    confidence = 85
                    # Check context
                    context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                    if any(term in context for term in ['fax', 'f:', 'fax:']):
                        priority = 'low'  # Fax numbers are lower risk
                        confidence = 70
                else:
                    confidence = 60
                    priority = 'low'
            
            # Email validation
            elif phi_type == 'Email':
                # Check if it's a patient email or institutional
                email_lower = value.lower()
                if any(domain in email_lower for domain in ['.gov', '.edu', '.org', 'hospital', 'clinic', 'medical']):
                    priority = 'low'
                    confidence = 60
                else:
                    if '@' in value and '.' in value.split('@')[1]:
                        confidence = 90
                        priority = 'medium'
            
            # Credit card validation with Luhn algorithm
            elif phi_type == 'Credit_Card':
                if validate_credit_card(value):
                    confidence = 95
                    priority = 'high'
                else:
                    continue  # Skip if not a valid credit card
            
            # Only add finding if confidence meets minimum threshold
            if confidence < 40:  # Lower threshold to 40% for edge cases
                continue
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': confidence,
                'priority': priority
            }
            
            findings.append(finding)
            found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def validate_credit_card(number: str) -> bool:
    """Validate credit card number using Luhn algorithm"""
    # Remove any non-digit characters
    number = re.sub(r'\D', '', number)
    
    # Check if it's a valid length
    if len(number) < 13 or len(number) > 19:
        return False
    
    # Luhn algorithm
    digits = [int(d) for d in number]
    checksum = 0
    
    # Double every second digit from the right
    for i in range(len(digits) - 2, -1, -2):
        doubled = digits[i] * 2
        if doubled > 9:
            doubled = doubled - 9
        digits[i] = doubled
    
    return sum(digits) % 10 == 0

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Special handling for patient names
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
            
            # Skip dates that are clearly not DOB (like years only)
            if phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})$')
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                except:
                    pass
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': calculate_confidence(phi_type, value, full_match),
                'priority': config.get('priority', 'medium')
            }
            
            # Only add if confidence is above threshold
            if finding['confidence'] >= 60:  # Minimum 60% confidence
                findings.append(finding)
                found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Initialize default values
            confidence = 70
            priority = config.get('priority', 'medium')
            skip_finding = False
            
            # Special handling for patient names with advanced context analysis
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
                
                # Analyze context to determine if it's really a patient name
                context_analysis = analyze_name_context(text, match.start(), match.end())
                
                if not context_analysis['is_patient']:
                    # Skip non-patient names unless confidence is very low
                    if context_analysis['confidence'] < 30:
                        continue
                
                confidence = context_analysis['confidence']
                priority = context_analysis['priority']
                
                # Add context type to description
                config['description'] = f"{config['description']} ({context_analysis['context_type']})"
            
            # Smart SSN detection
            elif phi_type == 'SSN':
                ssn_analysis = smart_ssn_detection(text, match)
                confidence = ssn_analysis['confidence']
                priority = ssn_analysis['priority']
            
            # Enhanced DOB detection
            elif phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Separate findings by actual risk level (using the smart priority from detection)
    high_risk_items = [f for f in findings if f.get('priority') == 'high' and f.get('confidence', 0) >= 70]
    medium_risk_items = [f for f in findings if f.get('priority') == 'medium' and f.get('confidence', 0) >= 60]
    low_risk_items = [f for f in findings if f.get('priority') == 'low']
    
    # Count unique PHI types at each risk level
    high_risk_types = set(f['type'] for f in high_risk_items)
    medium_risk_types = set(f['type'] for f in medium_risk_items)
    
    # Calculate base score with smart weighting
    high_risk_score = len(high_risk_items) * 15 + len(high_risk_types) * 10
    medium_risk_score = len(medium_risk_items) * 5 + len(medium_risk_types) * 5
    low_risk_score = len(low_risk_items) * 2
    
    # Bonus for particularly sensitive combinations
    sensitive_types = {'SSN', 'Credit Card', 'MRN'}
    if len(high_risk_types.intersection(sensitive_types)) >= 2:
        high_risk_score += 20
    
    # Check for identity theft risk (name + SSN or name + DOB + address)
    has_patient_name = any(f['type'] == 'Patient Name' and f.get('priority') == 'high' for f in findings)
    has_ssn = any(f['type'] == 'SSN' and f.get('confidence', 0) >= 80 for f in findings)
    has_dob = any(f['type'] == 'DOB' for f in findings)
    has_address = any(f['type'] == 'Address' for f in findings)
    
    if has_patient_name and has_ssn:
        high_risk_score += 25
    elif has_patient_name and has_dob and has_address:
        high_risk_score += 20
    
    # Total score calculation
    total_score = high_risk_score + medium_risk_score + low_risk_score
    
    # Determine risk level with more nuanced thresholds
    if len(high_risk_items) >= 3 or total_score >= 80 or (has_patient_name and has_ssn):
        return 'high', min(total_score, 95)
    elif len(high_risk_items) >= 1 or len(medium_risk_items) >= 3 or total_score >= 40:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Special handling for patient names
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
            
            # Skip dates that are clearly not DOB (like years only)
            if phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})$')
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                except:
                    pass
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': calculate_confidence(phi_type, value, full_match),
                'priority': config.get('priority', 'medium')
            }
            
            # Only add if confidence is above threshold
            if finding['confidence'] >= 60:  # Minimum 60% confidence
                findings.append(finding)
                found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
)
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                        
                        # Check context for DOB indicators
                        context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                        if any(term in context for term in ['dob', 'birth', 'born']):
                            confidence = 90
                            priority = 'medium'
                        else:
                            confidence = 50
                            priority = 'low'
                except:
                    confidence = 60
            
            # Enhanced MRN detection
            elif phi_type == 'MRN':
                context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                if any(term in context for term in ['mrn', 'medical record', 'mr#']):
                    confidence = 90
                    priority = 'high'
                elif any(term in context for term in ['account', 'acct']):
                    confidence = 75
                    priority = 'medium'
                else:
                    confidence = 50
                    priority = 'low'
            
            # Phone number validation
            elif phi_type == 'Phone':
                # Check if it's actually a phone number format
                phone_value = re.sub(r'[^\d]', '', value)
                if len(phone_value) == 10:
                    confidence = 85
                    # Check context
                    context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                    if any(term in context for term in ['fax', 'f:', 'fax:']):
                        priority = 'low'  # Fax numbers are lower risk
                        confidence = 70
                else:
                    confidence = 60
                    priority = 'low'
            
            # Email validation
            elif phi_type == 'Email':
                # Check if it's a patient email or institutional
                email_lower = value.lower()
                if any(domain in email_lower for domain in ['.gov', '.edu', '.org', 'hospital', 'clinic', 'medical']):
                    priority = 'low'
                    confidence = 60
                else:
                    if '@' in value and '.' in value.split('@')[1]:
                        confidence = 90
                        priority = 'medium'
            
            # Credit card validation with Luhn algorithm
            elif phi_type == 'Credit_Card':
                if validate_credit_card(value):
                    confidence = 95
                    priority = 'high'
                else:
                    continue  # Skip if not a valid credit card
            
            # Only add finding if confidence meets minimum threshold
            if confidence < 40:  # Lower threshold to 40% for edge cases
                continue
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': confidence,
                'priority': priority
            }
            
            findings.append(finding)
            found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def validate_credit_card(number: str) -> bool:
    """Validate credit card number using Luhn algorithm"""
    # Remove any non-digit characters
    number = re.sub(r'\D', '', number)
    
    # Check if it's a valid length
    if len(number) < 13 or len(number) > 19:
        return False
    
    # Luhn algorithm
    digits = [int(d) for d in number]
    checksum = 0
    
    # Double every second digit from the right
    for i in range(len(digits) - 2, -1, -2):
        doubled = digits[i] * 2
        if doubled > 9:
            doubled = doubled - 9
        digits[i] = doubled
    
    return sum(digits) % 10 == 0

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Special handling for patient names
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
            
            # Skip dates that are clearly not DOB (like years only)
            if phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})$')
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                except:
                    pass
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': calculate_confidence(phi_type, value, full_match),
                'priority': config.get('priority', 'medium')
            }
            
            # Only add if confidence is above threshold
            if finding['confidence'] >= 60:  # Minimum 60% confidence
                findings.append(finding)
                found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
)
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                        
                        # Check context for DOB indicators
                        context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                        if any(term in context for term in ['dob', 'birth', 'born']):
                            confidence = 90
                            priority = 'medium'
                        else:
                            confidence = 50
                            priority = 'low'
                except:
                    confidence = 60

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Separate findings by actual risk level (using the smart priority from detection)
    high_risk_items = [f for f in findings if f.get('priority') == 'high' and f.get('confidence', 0) >= 70]
    medium_risk_items = [f for f in findings if f.get('priority') == 'medium' and f.get('confidence', 0) >= 60]
    low_risk_items = [f for f in findings if f.get('priority') == 'low']
    
    # Count unique PHI types at each risk level
    high_risk_types = set(f['type'] for f in high_risk_items)
    medium_risk_types = set(f['type'] for f in medium_risk_items)
    
    # Calculate base score with smart weighting
    high_risk_score = len(high_risk_items) * 15 + len(high_risk_types) * 10
    medium_risk_score = len(medium_risk_items) * 5 + len(medium_risk_types) * 5
    low_risk_score = len(low_risk_items) * 2
    
    # Bonus for particularly sensitive combinations
    sensitive_types = {'SSN', 'Credit Card', 'MRN'}
    if len(high_risk_types.intersection(sensitive_types)) >= 2:
        high_risk_score += 20
    
    # Check for identity theft risk (name + SSN or name + DOB + address)
    has_patient_name = any(f['type'] == 'Patient Name' and f.get('priority') == 'high' for f in findings)
    has_ssn = any(f['type'] == 'SSN' and f.get('confidence', 0) >= 80 for f in findings)
    has_dob = any(f['type'] == 'DOB' for f in findings)
    has_address = any(f['type'] == 'Address' for f in findings)
    
    if has_patient_name and has_ssn:
        high_risk_score += 25
    elif has_patient_name and has_dob and has_address:
        high_risk_score += 20
    
    # Total score calculation
    total_score = high_risk_score + medium_risk_score + low_risk_score
    
    # Determine risk level with more nuanced thresholds
    if len(high_risk_items) >= 3 or total_score >= 80 or (has_patient_name and has_ssn):
        return 'high', min(total_score, 95)
    elif len(high_risk_items) >= 1 or len(medium_risk_items) >= 3 or total_score >= 40:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Special handling for patient names
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
            
            # Skip dates that are clearly not DOB (like years only)
            if phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})$')
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                except:
                    pass
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': calculate_confidence(phi_type, value, full_match),
                'priority': config.get('priority', 'medium')
            }
            
            # Only add if confidence is above threshold
            if finding['confidence'] >= 60:  # Minimum 60% confidence
                findings.append(finding)
                found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
)
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                        
                        # Check context for DOB indicators
                        context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                        if any(term in context for term in ['dob', 'birth', 'born']):
                            confidence = 90
                            priority = 'medium'
                        else:
                            confidence = 50
                            priority = 'low'
                except:
                    confidence = 60
            
            # Enhanced MRN detection
            elif phi_type == 'MRN':
                context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                if any(term in context for term in ['mrn', 'medical record', 'mr#']):
                    confidence = 90
                    priority = 'high'
                elif any(term in context for term in ['account', 'acct']):
                    confidence = 75
                    priority = 'medium'
                else:
                    confidence = 50
                    priority = 'low'
            
            # Phone number validation
            elif phi_type == 'Phone':
                # Check if it's actually a phone number format
                phone_value = re.sub(r'[^\d]', '', value)
                if len(phone_value) == 10:
                    confidence = 85
                    # Check context
                    context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                    if any(term in context for term in ['fax', 'f:', 'fax:']):
                        priority = 'low'  # Fax numbers are lower risk
                        confidence = 70
                else:
                    confidence = 60
                    priority = 'low'
            
            # Email validation
            elif phi_type == 'Email':
                # Check if it's a patient email or institutional
                email_lower = value.lower()
                if any(domain in email_lower for domain in ['.gov', '.edu', '.org', 'hospital', 'clinic', 'medical']):
                    priority = 'low'
                    confidence = 60
                else:
                    if '@' in value and '.' in value.split('@')[1]:
                        confidence = 90
                        priority = 'medium'
            
            # Credit card validation with Luhn algorithm
            elif phi_type == 'Credit_Card':
                if validate_credit_card(value):
                    confidence = 95
                    priority = 'high'
                else:
                    continue  # Skip if not a valid credit card
            
            # Only add finding if confidence meets minimum threshold
            if confidence < 40:  # Lower threshold to 40% for edge cases
                continue
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': confidence,
                'priority': priority
            }
            
            findings.append(finding)
            found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def validate_credit_card(number: str) -> bool:
    """Validate credit card number using Luhn algorithm"""
    # Remove any non-digit characters
    number = re.sub(r'\D', '', number)
    
    # Check if it's a valid length
    if len(number) < 13 or len(number) > 19:
        return False
    
    # Luhn algorithm
    digits = [int(d) for d in number]
    checksum = 0
    
    # Double every second digit from the right
    for i in range(len(digits) - 2, -1, -2):
        doubled = digits[i] * 2
        if doubled > 9:
            doubled = doubled - 9
        digits[i] = doubled
    
    return sum(digits) % 10 == 0

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Special handling for patient names
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
            
            # Skip dates that are clearly not DOB (like years only)
            if phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})$')
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                except:
                    pass
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': calculate_confidence(phi_type, value, full_match),
                'priority': config.get('priority', 'medium')
            }
            
            # Only add if confidence is above threshold
            if finding['confidence'] >= 60:  # Minimum 60% confidence
                findings.append(finding)
                found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Initialize default values
            confidence = 70
            priority = config.get('priority', 'medium')
            skip_finding = False
            
            # Special handling for patient names with advanced context analysis
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
                
                # Analyze context to determine if it's really a patient name
                context_analysis = analyze_name_context(text, match.start(), match.end())
                
                if not context_analysis['is_patient']:
                    # Skip non-patient names unless confidence is very low
                    if context_analysis['confidence'] < 30:
                        continue
                
                confidence = context_analysis['confidence']
                priority = context_analysis['priority']
                
                # Add context type to description
                config['description'] = f"{config['description']} ({context_analysis['context_type']})"
            
            # Smart SSN detection
            elif phi_type == 'SSN':
                ssn_analysis = smart_ssn_detection(text, match)
                confidence = ssn_analysis['confidence']
                priority = ssn_analysis['priority']
            
            # Enhanced DOB detection
            elif phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Separate findings by actual risk level (using the smart priority from detection)
    high_risk_items = [f for f in findings if f.get('priority') == 'high' and f.get('confidence', 0) >= 70]
    medium_risk_items = [f for f in findings if f.get('priority') == 'medium' and f.get('confidence', 0) >= 60]
    low_risk_items = [f for f in findings if f.get('priority') == 'low']
    
    # Count unique PHI types at each risk level
    high_risk_types = set(f['type'] for f in high_risk_items)
    medium_risk_types = set(f['type'] for f in medium_risk_items)
    
    # Calculate base score with smart weighting
    high_risk_score = len(high_risk_items) * 15 + len(high_risk_types) * 10
    medium_risk_score = len(medium_risk_items) * 5 + len(medium_risk_types) * 5
    low_risk_score = len(low_risk_items) * 2
    
    # Bonus for particularly sensitive combinations
    sensitive_types = {'SSN', 'Credit Card', 'MRN'}
    if len(high_risk_types.intersection(sensitive_types)) >= 2:
        high_risk_score += 20
    
    # Check for identity theft risk (name + SSN or name + DOB + address)
    has_patient_name = any(f['type'] == 'Patient Name' and f.get('priority') == 'high' for f in findings)
    has_ssn = any(f['type'] == 'SSN' and f.get('confidence', 0) >= 80 for f in findings)
    has_dob = any(f['type'] == 'DOB' for f in findings)
    has_address = any(f['type'] == 'Address' for f in findings)
    
    if has_patient_name and has_ssn:
        high_risk_score += 25
    elif has_patient_name and has_dob and has_address:
        high_risk_score += 20
    
    # Total score calculation
    total_score = high_risk_score + medium_risk_score + low_risk_score
    
    # Determine risk level with more nuanced thresholds
    if len(high_risk_items) >= 3 or total_score >= 80 or (has_patient_name and has_ssn):
        return 'high', min(total_score, 95)
    elif len(high_risk_items) >= 1 or len(medium_risk_items) >= 3 or total_score >= 40:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Special handling for patient names
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
            
            # Skip dates that are clearly not DOB (like years only)
            if phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})$')
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                except:
                    pass
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': calculate_confidence(phi_type, value, full_match),
                'priority': config.get('priority', 'medium')
            }
            
            # Only add if confidence is above threshold
            if finding['confidence'] >= 60:  # Minimum 60% confidence
                findings.append(finding)
                found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
)
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                        
                        # Check context for DOB indicators
                        context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                        if any(term in context for term in ['dob', 'birth', 'born']):
                            confidence = 90
                            priority = 'medium'
                        else:
                            confidence = 50
                            priority = 'low'
                except:
                    confidence = 60
            
            # Enhanced MRN detection
            elif phi_type == 'MRN':
                context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                if any(term in context for term in ['mrn', 'medical record', 'mr#']):
                    confidence = 90
                    priority = 'high'
                elif any(term in context for term in ['account', 'acct']):
                    confidence = 75
                    priority = 'medium'
                else:
                    confidence = 50
                    priority = 'low'
            
            # Phone number validation
            elif phi_type == 'Phone':
                # Check if it's actually a phone number format
                phone_value = re.sub(r'[^\d]', '', value)
                if len(phone_value) == 10:
                    confidence = 85
                    # Check context
                    context = text[max(0, match.start()-30):min(len(text), match.end()+30)].lower()
                    if any(term in context for term in ['fax', 'f:', 'fax:']):
                        priority = 'low'  # Fax numbers are lower risk
                        confidence = 70
                else:
                    confidence = 60
                    priority = 'low'
            
            # Email validation
            elif phi_type == 'Email':
                # Check if it's a patient email or institutional
                email_lower = value.lower()
                if any(domain in email_lower for domain in ['.gov', '.edu', '.org', 'hospital', 'clinic', 'medical']):
                    priority = 'low'
                    confidence = 60
                else:
                    if '@' in value and '.' in value.split('@')[1]:
                        confidence = 90
                        priority = 'medium'
            
            # Credit card validation with Luhn algorithm
            elif phi_type == 'Credit_Card':
                if validate_credit_card(value):
                    confidence = 95
                    priority = 'high'
                else:
                    continue  # Skip if not a valid credit card
            
            # Only add finding if confidence meets minimum threshold
            if confidence < 40:  # Lower threshold to 40% for edge cases
                continue
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': confidence,
                'priority': priority
            }
            
            findings.append(finding)
            found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def validate_credit_card(number: str) -> bool:
    """Validate credit card number using Luhn algorithm"""
    # Remove any non-digit characters
    number = re.sub(r'\D', '', number)
    
    # Check if it's a valid length
    if len(number) < 13 or len(number) > 19:
        return False
    
    # Luhn algorithm
    digits = [int(d) for d in number]
    checksum = 0
    
    # Double every second digit from the right
    for i in range(len(digits) - 2, -1, -2):
        doubled = digits[i] * 2
        if doubled > 9:
            doubled = doubled - 9
        digits[i] = doubled
    
    return sum(digits) % 10 == 0

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)
, value):
        confidence = 85
        priority = 'high'
    
    return {
        'confidence': confidence,
        'priority': priority
    }

def detect_phi(text: str) -> List[Dict]:
    """Detect PHI in text using advanced regex patterns with context awareness"""
    findings = []
    found_positions = set()  # Track positions to avoid duplicates
    
    for phi_type, config in PHI_PATTERNS.items():
        matches = re.finditer(config['pattern'], text, re.IGNORECASE | re.MULTILINE)
        for match in matches:
            # Skip if we've already found PHI at this position
            if match.start() in found_positions:
                continue
                
            # Get the full match and the captured group (if any)
            full_match = match.group(0)
            value = match.group(1) if match.groups() else match.group(0)
            
            # Special handling for patient names
            if phi_type == 'Patient_Name':
                # Extract just the name part, not the label
                value = match.group(1) if match.groups() else value
                # Skip if it's a medical term
                if is_medical_term(value):
                    continue
            
            # Skip dates that are clearly not DOB (like years only)
            if phi_type == 'DOB':
                # Extract year using a simpler approach
                try:
                    # Look for 4-digit or 2-digit year at the end
                    year_pattern = re.compile(r'(\d{4}|\d{2})$')
                    year_match = year_pattern.search(value)
                    if year_match:
                        year = int(year_match.group(1))
                        # Convert 2-digit year to 4-digit
                        if year < 100:
                            year = 1900 + year if year > 50 else 2000 + year
                        # Skip if date is too old (before 1900) or in the future
                        current_year = datetime.now().year
                        if year < 1900 or year > current_year:
                            continue
                except:
                    pass
            
            # Redact the value appropriately
            if len(value) > 4:
                if phi_type == 'Email':
                    # Keep first letter and domain
                    parts = value.split('@')
                    if len(parts) == 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + '@' + parts[1]
                    else:
                        redacted = value[:1] + '*' * (len(value) - 3) + value[-2:]
                elif phi_type == 'Patient_Name':
                    # Show first letter of first name and last initial
                    parts = value.split()
                    if len(parts) >= 2:
                        redacted = parts[0][0] + '*' * (len(parts[0]) - 1) + ' ' + parts[-1][0] + '*' * (len(parts[-1]) - 1)
                    else:
                        redacted = value[0] + '*' * (len(value) - 1)
                else:
                    redacted = value[:2] + '*' * (len(value) - 4) + value[-2:]
            else:
                redacted = '*' * len(value)
            
            finding = {
                'type': phi_type.replace('_', ' '),
                'description': config['description'],
                'value': redacted,
                'position': match.span(),
                'confidence': calculate_confidence(phi_type, value, full_match),
                'priority': config.get('priority', 'medium')
            }
            
            # Only add if confidence is above threshold
            if finding['confidence'] >= 60:  # Minimum 60% confidence
                findings.append(finding)
                found_positions.add(match.start())
    
    # Remove duplicates and sort by position
    findings = sorted(findings, key=lambda x: x['position'][0])
    
    return findings

def calculate_confidence(phi_type: str, value: str, full_match: str = "") -> int:
    """Calculate confidence score for PHI detection with context awareness"""
    base_confidence = 70
    
    # High confidence for explicit labels
    if any(label in full_match.lower() for label in ['patient name:', 'dob:', 'ssn:', 'mrn:']):
        return 95
    
    # Specific pattern confidence adjustments
    if phi_type == 'SSN':
        if re.match(r'^\d{3}-\d{2}-\d{4}$', value):
            return 95
        elif re.match(r'^\d{9}$', value):
            return 85
    
    elif phi_type == 'MRN':
        if 'MRN' in full_match or 'Medical Record' in full_match:
            return 90
        return 75
    
    elif phi_type == 'Email':
        if '@' in value and '.' in value.split('@')[1]:
            return 90
    
    elif phi_type == 'Patient_Name':
        # Higher confidence if preceded by patient-related terms
        if any(term in full_match.lower() for term in ['patient name', 'patient:', 'name:']):
            return 90
        # Lower confidence for generic names
        return 65
    
    elif phi_type == 'DOB':
        # Check if it looks like a realistic birth date
        try:
            # Simple year extraction
            year = int(re.search(r'\d{4}|\d{2}$', value).group())
            if year < 100:
                year = 1900 + year if year > 50 else 2000 + year
            current_year = datetime.now().year
            age = current_year - year
            if 0 <= age <= 120:  # Realistic age range
                return 85
        except:
            pass
        return 70
    
    elif phi_type == 'Phone':
        if re.match(r'^\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}$', value):
            return 85
    
    elif phi_type == 'Address':
        # Higher confidence for complete addresses
        if any(apt in full_match.lower() for apt in ['apt', 'suite', 'unit']):
            return 85
        return 75
    
    elif phi_type == 'Credit_Card':
        # Luhn algorithm check could be added here
        return 90
    
    return base_confidence

def calculate_risk_score(findings: List[Dict]) -> Tuple[str, int]:
    """Calculate document risk score based on PHI findings with smart weighting"""
    if not findings:
        return 'low', 0
    
    # Count unique PHI types, not total instances
    phi_types_found = set()
    high_risk_items = 0
    
    for finding in findings:
        phi_type = finding['type']
        phi_types_found.add(phi_type)
        
        # Count high-risk items (SSN, Credit Card, etc.)
        if finding.get('priority') == 'high' and phi_type in ['SSN', 'Credit Card', 'MRN']:
            high_risk_items += 1
    
    # Base score on diversity of PHI types found
    base_score = len(phi_types_found) * 15
    
    # Add points for high-risk items
    high_risk_score = high_risk_items * 20
    
    # Add small penalty for volume (but not too much)
    volume_score = min(len(findings) * 2, 20)  # Max 20 points for volume
    
    # Total score
    total_score = base_score + high_risk_score + volume_score
    
    # More reasonable risk levels
    if total_score >= 80 or high_risk_items >= 2:
        return 'high', min(total_score, 95)  # Cap at 95 instead of 100
    elif total_score >= 40 or high_risk_items >= 1:
        return 'moderate', total_score
    else:
        return 'low', total_score

def extract_text_from_file(file) -> str:
    """Extract text from various file types"""
    text = ""
    
    try:
        if file.type == "application/pdf":
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n"
                
        elif file.type == "text/plain":
            # Read text file
            text = str(file.read(), 'utf-8')
            
        elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                          "application/msword"] and docx:
            # Read Word document
            doc = docx.Document(file)
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
        else:
            st.warning(f"File type {file.type} not fully supported. Using basic text extraction.")
            try:
                text = str(file.read(), 'utf-8', errors='ignore')
            except:
                text = ""
            
    except Exception as e:
        st.error(f"Error reading {file.name}: {str(e)}")
        return ""
    
    return text

def process_file(file, min_confidence: int = 60) -> Dict:
    """Process uploaded file and scan for PHI"""
    # Extract actual text from the file
    text = extract_text_from_file(file)
    
    if not text:
        st.error(f"Could not extract text from {file.name}")
        return {
            'filename': file.name,
            'file_size': f"{file.size / 1024:.2f} KB",
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'findings': [],
            'risk_level': 'low',
            'risk_score': 0,
            'text_preview': "No text extracted"
        }
    
    # Detect PHI in the actual document text
    all_findings = detect_phi(text)
    
    # Filter by confidence threshold
    findings = [f for f in all_findings if f['confidence'] >= min_confidence]
    
    # Calculate risk score
    risk_level, risk_score = calculate_risk_score(findings)
    
    # Create a preview of the text (first 500 characters)
    text_preview = text[:500] + '...' if len(text) > 500 else text
    
    return {
        'filename': file.name,
        'file_size': f"{file.size / 1024:.2f} KB",
        'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'findings': findings,
        'risk_level': risk_level,
        'risk_score': risk_score,
        'text_preview': text_preview,
        'full_text': text  # Store full text for redaction
    }

def redact_text(text: str, findings: List[Dict]) -> str:
    """Redact PHI from text"""
    # Sort findings by position (reverse order to maintain positions)
    sorted_findings = sorted(findings, key=lambda x: x['position'][0], reverse=True)
    
    redacted_text = text
    for finding in sorted_findings:
        start, end = finding['position']
        redacted_text = redacted_text[:start] + '[REDACTED]' + redacted_text[end:]
    
    return redacted_text

def generate_report_content(result: Dict) -> str:
    """Generate a detailed PHI scan report"""
    report = f"""# PHI Scan Report

## Document Information
- **Filename:** {result['filename']}
- **File Size:** {result['file_size']}
- **Scan Date:** {result['timestamp']}
- **Risk Level:** {result['risk_level'].upper()}
- **Risk Score:** {result['risk_score']}/100

## Executive Summary
This document was scanned for Protected Health Information (PHI) as defined by HIPAA regulations. 
The scan identified **{len(result['findings'])}** potential PHI element(s) with a risk assessment of **{result['risk_level'].upper()}**.

## PHI Findings Details
"""
    
    if result['findings']:
        # Group findings by type
        findings_by_type = {}
        for finding in result['findings']:
            phi_type = finding['type']
            if phi_type not in findings_by_type:
                findings_by_type[phi_type] = []
            findings_by_type[phi_type].append(finding)
        
        # Report each type
        for phi_type, findings in findings_by_type.items():
            report += f"\n### {phi_type} ({len(findings)} instance(s))\n"
            report += f"- **Risk Priority:** {findings[0].get('priority', 'medium').upper()}\n"
            report += f"- **Instances Found:**\n"
            for i, finding in enumerate(findings, 1):
                report += f"  {i}. {finding['value']} (Confidence: {finding['confidence']}%)\n"
    else:
        report += "\n*No PHI elements were detected in this document.*\n"
    
    # Add recommendations
    report += "\n## Recommendations\n"
    if result['risk_level'] == 'high':
        report += """- **IMMEDIATE ACTION REQUIRED**: This document contains high-risk PHI elements.
- Ensure proper encryption when storing or transmitting this document.
- Limit access to authorized personnel only.
- Consider redacting sensitive information before sharing.
- Log all access to this document for HIPAA compliance."""
    elif result['risk_level'] == 'moderate':
        report += """- This document contains moderate-risk PHI elements.
- Apply standard PHI handling procedures.
- Ensure secure transmission channels are used.
- Verify recipient authorization before sharing."""
    else:
        report += """- This document contains low-risk PHI elements.
- Follow standard privacy protocols.
- Maintain audit trails as per organizational policy."""
    
    # Add compliance notes
    report += """\n\n## Compliance Notes
- This scan was performed using pattern-based detection algorithms.
- Results should be reviewed by qualified personnel for accuracy.
- This report is for internal use only and contains sensitive information.
- Retain this report as per your organization's HIPAA documentation requirements.

---
*Generated by Sentinel PHI Scanner v1.0*"""
    
    return report

# Main App Layout
st.title("🔐 Sentinel: AI PHI Risk Scanner")
st.markdown("**Protect patient privacy with intelligent PHI detection and risk assessment**")

# Sidebar
with st.sidebar:
    st.header("📊 Dashboard Stats")
    
    col1, col2 = st.columns(2)
    with col1:
        st.metric("Files Scanned", st.session_state.total_files_scanned)
    with col2:
        st.metric("High Risk", st.session_state.high_risk_count, 
                 delta=None if st.session_state.high_risk_count == 0 else "⚠️")
    
    st.metric("Total PHI Found", st.session_state.total_phi_found)
    
    st.divider()
    
    st.header("⚙️ Settings")
    
    # Scan sensitivity
    sensitivity = st.select_slider(
        "Detection Sensitivity",
        options=['Low', 'Medium', 'High'],
        value='Medium',
        help="Low: 80%+ confidence, Medium: 60%+ confidence, High: 40%+ confidence"
    )
    
    # Confidence threshold based on sensitivity
    confidence_thresholds = {'Low': 80, 'Medium': 60, 'High': 40}
    min_confidence = confidence_thresholds[sensitivity]
    
    # Auto-redact option
    auto_redact = st.checkbox("Auto-redact PHI in reports", value=False)
    
    # Show confidence scores
    show_confidence = st.checkbox("Show confidence scores", value=True)
    
    # Export format
    export_format = st.selectbox(
        "Export Format",
        ['PDF', 'CSV', 'JSON']
    )

# Main content area - Store active tab in session state
if 'active_tab' not in st.session_state:
    st.session_state.active_tab = "📤 Upload & Scan"

# Create tabs
tab_names = ["📤 Upload & Scan", "📋 Scan History", "📊 Analytics"]
tab1, tab2, tab3 = st.tabs(tab_names)

with tab1:
    # Check if we should show results or upload
    if st.session_state.current_results:
        # Results View
        st.header("📄 Scan Results")
        
        # Action buttons at the top
        col1, col2 = st.columns([1, 5])
        with col1:
            if st.button("🔙 Back to Upload", key="back_to_upload"):
                st.session_state.current_results = []
                st.rerun()
        
        with col2:
            if st.button("🗑️ Clear All Results", key="clear_all_results"):
                st.session_state.current_results = []
                st.rerun()
        
        # Summary Statistics
        high_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'high']
        moderate_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'moderate']
        low_risk_files = [r for r in st.session_state.current_results if r['risk_level'] == 'low']
        
        # Alert Box for High Risk Files
        if high_risk_files:
            st.error(f"""
            🚨 **ATTENTION: {len(high_risk_files)} HIGH-RISK FILE(S) DETECTED**
            
            These files contain sensitive PHI and require immediate security measures. 
            Review handling recommendations below and notify your compliance officer.
            """)
        
        # Summary Metrics
        st.markdown("### 📊 Scan Summary")
        sum_col1, sum_col2, sum_col3, sum_col4 = st.columns(4)
        
        with sum_col1:
            st.metric(
                "Total Files Scanned", 
                len(st.session_state.current_results),
                delta=None
            )
        
        with sum_col2:
            st.metric(
                "High Risk", 
                len(high_risk_files),
                delta=None if not high_risk_files else "⚠️",
                delta_color="inverse"
            )
        
        with sum_col3:
            st.metric(
                "Moderate Risk", 
                len(moderate_risk_files),
                delta=None
            )
        
        with sum_col4:
            st.metric(
                "Low Risk", 
                len(low_risk_files),
                delta=None
            )
        
        st.divider()
        
        # Display results
        for idx, result in enumerate(st.session_state.current_results):
            # Create unique container for each result
            with st.container():
                # Result header
                st.subheader(f"📄 {result['filename']}")
                
                # Prominent Risk Assessment Box
                risk_level = result['risk_level']
                risk_configs = {
                    'high': {
                        'color': '#ff4757',
                        'bg_color': '#ffebee',
                        'icon': '🚨',
                        'label': 'HIGH RISK',
                        'action': 'IMMEDIATE ACTION REQUIRED',
                        'recommendations': [
                            '🔒 **Encrypt immediately** - This file must be encrypted at rest and in transit',
                            '👤 **Restrict access** - Limit to authorized personnel only',
                            '📝 **Log all access** - Document who accesses this file and when',
                            '🔐 **Consider redaction** - Remove PHI before sharing externally',
                            '⚡ **Report to compliance** - Notify your HIPAA compliance officer'
                        ]
                    },
                    'moderate': {
                        'color': '#ff9f43',
                        'bg_color': '#fff8e1',
                        'icon': '⚠️',
                        'label': 'MODERATE RISK',
                        'action': 'STANDARD PHI PROTOCOLS APPLY',
                        'recommendations': [
                            '🔒 **Apply standard encryption** - Use organizational encryption standards',
                            '✅ **Verify recipient authorization** - Confirm BAA is in place',
                            '📧 **Use secure channels** - Send via encrypted email or secure portal',
                            '📋 **Follow PHI procedures** - Apply your organization\'s standard protocols',
                            '🗂️ **Maintain audit trail** - Keep records of file handling'
                        ]
                    },
                    'low': {
                        'color': '#10ac84',
                        'bg_color': '#e8f5e9',
                        'icon': '✅',
                        'label': 'LOW RISK',
                        'action': 'MINIMAL PHI DETECTED',
                        'recommendations': [
                            '✅ **Follow basic protocols** - Standard privacy measures apply',
                            '🔍 **Review before sharing** - Quick check before external distribution',
                            '📁 **Store securely** - Use standard secure storage',
                            '📝 **Document as needed** - Follow organizational policies',
                            '👍 **Safe for most uses** - Minimal PHI exposure risk'
                        ]
                    }
                }
                
                config = risk_configs[risk_level]
                
                # Risk Assessment Card
                st.markdown(f"""
                <div style="
                    background-color: {config['bg_color']};
                    border-left: 5px solid {config['color']};
                    padding: 20px;
                    border-radius: 10px;
                    margin: 20px 0;
                ">
                    <div style="display: flex; align-items: center; justify-content: space-between;">
                        <div>
                            <h2 style="color: {config['color']}; margin: 0;">
                                {config['icon']} {config['label']}
                            </h2>
                            <p style="font-size: 16px; font-weight: bold; margin: 10px 0 0 0;">
                                {config['action']}
                            </p>
                        </div>
                        <div style="text-align: right;">
                            <div style="font-size: 48px; font-weight: bold; color: {config['color']};">
                                {result['risk_score']}
                            </div>
                            <div style="font-size: 14px; color: #666;">Risk Score</div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Handling Recommendations
                with st.expander("📋 **Handling Recommendations**", expanded=risk_level in ['high', 'moderate']):
                    st.markdown("### How to Handle This File:")
                    for rec in config['recommendations']:
                        st.markdown(rec)
                
                # File info in a cleaner layout
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("File Size", result['file_size'])
                with col2:
                    st.metric("PHI Elements", len(result['findings']))
                with col3:
                    phi_types = len(set(f['type'] for f in result['findings'])) if result['findings'] else 0
                    st.metric("PHI Types", phi_types)
                with col4:
                    high_priority = len([f for f in result['findings'] if f.get('priority') == 'high'])
                    st.metric("High Priority", high_priority)
                
                # PHI Findings
                if result['findings']:
                    st.markdown("### 🔍 PHI Elements Detected")
                    
                    # Create findings dataframe
                    findings_df = pd.DataFrame(result['findings'])
                    findings_df = findings_df[['type', 'description', 'value', 'confidence']]
                    findings_df.columns = ['Type', 'Description', 'Redacted Value', 'Confidence %']
                    
                    st.dataframe(findings_df, use_container_width=True, hide_index=True)
                    
                    # Export Options in a clean layout
                    st.markdown("### 📥 Export Options")
                    
                    # Generate report content once
                    report_content = generate_report_content(result)
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        # Markdown Report
                        report_bytes = report_content.encode('utf-8')
                        report_filename = f"PHI_Report_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
                        
                        st.download_button(
                            label="📄 Markdown Report",
                            data=report_bytes,
                            file_name=report_filename,
                            mime="text/markdown",
                            key=f"md_{idx}_{result['filename']}"
                        )
                    
                    with col2:
                        # CSV
                        csv_data = findings_df.to_csv(index=False)
                        csv_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
                        
                        st.download_button(
                            label="📊 CSV Data",
                            data=csv_data,
                            file_name=csv_filename,
                            mime="text/csv",
                            key=f"csv_{idx}_{result['filename']}"
                        )
                    
                    with col3:
                        # JSON
                        json_data = json.dumps(result['findings'], indent=2)
                        json_filename = f"PHI_Findings_{result['filename'].split('.')[0]}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                        
                        st.download_button(
                            label="🗂️ JSON Data",
                            data=json_data,
                            file_name=json_filename,
                            mime="application/json",
                            key=f"json_{idx}_{result['filename']}"
                        )
                    
                    with col4:
                        # Redacted Document
                        if 'full_text' in result:
                            redacted = redact_text(result['full_text'], result['findings'])
                            redacted_filename = f"REDACTED_{result['filename']}"
                            
                            st.download_button(
                                label="🔒 Redacted Doc",
                                data=redacted,
                                file_name=redacted_filename,
                                mime="text/plain",
                                key=f"redacted_{idx}_{result['filename']}"
                            )
                
                else:
                    st.success("✅ No PHI detected in this document")
                
                st.divider()
    
    else:
        # Upload View
        st.header("Upload Documents for PHI Scanning")
        
        uploaded_files = st.file_uploader(
            "Drop files here or click to upload",
            type=['txt', 'pdf', 'doc', 'docx'],
            accept_multiple_files=True,
            help="Supports PDF, TXT, DOC files • Simulates fax intake"
        )
        
        if uploaded_files:
            # Process button
            if st.button("🔍 Scan for PHI", type="primary"):
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                results = []
                for idx, file in enumerate(uploaded_files):
                    status_text.text(f"Scanning {file.name}...")
                    progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    # Process file
                    result = process_file(file, min_confidence)
                    results.append(result)
                    
                    # Update session state
                    st.session_state.scan_history.append(result)
                    st.session_state.total_files_scanned += 1
                    if result['risk_level'] == 'high':
                        st.session_state.high_risk_count += 1
                    st.session_state.total_phi_found += len(result['findings'])
                    
                    time.sleep(0.5)  # Simulate processing time
                
                progress_bar.empty()
                status_text.empty()
                
                # Store results and trigger rerun to show them
                st.session_state.current_results = results
                st.rerun()

with tab2:
    st.header("📋 Scan History")
    
    if st.session_state.scan_history:
        # Create history dataframe
        history_data = []
        for scan in st.session_state.scan_history:
            history_data.append({
                'Timestamp': scan['timestamp'],
                'Filename': scan['filename'],
                'Risk Level': scan['risk_level'].upper(),
                'PHI Found': len(scan['findings']),
                'Risk Score': scan['risk_score']
            })
        
        history_df = pd.DataFrame(history_data)
        
        # Filter options
        col1, col2 = st.columns([1, 3])
        with col1:
            risk_filter = st.multiselect(
                "Filter by Risk Level",
                ['HIGH', 'MODERATE', 'LOW'],
                default=['HIGH', 'MODERATE', 'LOW']
            )
        
        # Apply filters
        filtered_df = history_df[history_df['Risk Level'].isin(risk_filter)]
        
        # Display filtered history
        st.dataframe(
            filtered_df,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Risk Level": st.column_config.TextColumn(
                    "Risk Level",
                    help="Risk assessment based on PHI findings"
                ),
                "Risk Score": st.column_config.ProgressColumn(
                    "Risk Score",
                    help="Overall risk score (0-100)",
                    format="%d",
                    min_value=0,
                    max_value=100,
                ),
            }
        )
        
        # Export history button
        if st.button("📥 Export Full History"):
            csv = filtered_df.to_csv(index=False)
            st.download_button(
                label="Download History CSV",
                data=csv,
                file_name=f"phi_scan_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                mime="text/csv"
            )
    else:
        st.info("No scan history yet. Upload and scan some documents to get started!")

with tab3:
    st.header("📊 PHI Detection Analytics")
    
    if st.session_state.scan_history:
        # Risk distribution
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("Risk Level Distribution")
            risk_counts = pd.DataFrame(st.session_state.scan_history)['risk_level'].value_counts()
            st.bar_chart(risk_counts)
        
        with col2:
            st.subheader("PHI Types Detected")
            all_types = []
            for scan in st.session_state.scan_history:
                all_types.extend([f['type'] for f in scan['findings']])
            
            if all_types:
                type_counts = pd.Series(all_types).value_counts()
                st.bar_chart(type_counts)
            else:
                st.info("No PHI detected yet")
        
        # Trend over time
        st.subheader("Detection Trends")
        if len(st.session_state.scan_history) > 1:
            trend_data = []
            for scan in st.session_state.scan_history:
                trend_data.append({
                    'Timestamp': pd.to_datetime(scan['timestamp']),
                    'PHI Count': len(scan['findings']),
                    'Risk Score': scan['risk_score']
                })
            
            trend_df = pd.DataFrame(trend_data)
            trend_df.set_index('Timestamp', inplace=True)
            
            st.line_chart(trend_df)
    else:
        st.info("No data available yet. Start scanning documents to see analytics!")

# Footer
st.divider()
st.markdown("""
<center>
    <small>
    🔐 Sentinel PHI Scanner v1.0 | Built for HIIM Professionals<br>
    <em>Remember: This is a demonstration tool. Always follow your organization's PHI handling policies.</em>
    </small>
</center>
""", unsafe_allow_html=True)